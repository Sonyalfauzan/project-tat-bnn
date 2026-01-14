"""
=================================================================================
SISTEM ASESMEN TERPADU (TAT) BNN PROVINSI KALIMANTAN UTARA
=================================================================================
Berdasarkan:
- UU No. 35 Tahun 2009 tentang Narkotika
- KEP/99 I/X/KA/PB/06.00/2025/BNN tentang Petunjuk Teknis Pelaksanaan 
  Asesmen Terpadu
- Instrumen: ASI, ASAM, DSM-5, ICD-10, PPDGJ III
=================================================================================
"""

import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
import json

# ReportLab untuk PDF
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# python-docx untuk Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# KONFIGURASI
# =============================================================================
st.set_page_config(
    page_title="TAT BNN Kalimantan Utara",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1.5rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .info-box {
        background-color: #d1ecf1;
        border-left: 4px solid #17a2b8;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border-left: 4px solid #ffc107;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# KONSTANTA
# =============================================================================

JENIS_NARKOTIKA_LIST = [
    "Heroin", "Ganja", "Ekstasi/MDMA", "Sabu/Metamfetamin", 
    "Kokain", "Carisoprodol", "Cannabinoid Sintesis", 
    "Morfin", "Kodein", "Lainnya"
]

DIAGNOSIS_ICD10 = {
    "F11": "Gangguan Mental dan Perilaku akibat Penggunaan Opioid",
    "F12": "Gangguan Mental dan Perilaku akibat Penggunaan Cannabinoid",
    "F14": "Gangguan Mental dan Perilaku akibat Penggunaan Kokain",
    "F15": "Gangguan Mental dan Perilaku akibat Penggunaan Stimulan (Amfetamin, Metamfetamin)",
    "F19": "Gangguan Mental dan Perilaku akibat Penggunaan Zat Multipel"
}

DSM5_CRITERIA = [
    "Menggunakan dalam jumlah/waktu lebih lama dari yang direncanakan",
    "Keinginan kuat/gagal mengurangi penggunaan",
    "Banyak waktu untuk mendapatkan/menggunakan/pulih dari efek",
    "Craving (keinginan kuat menggunakan)",
    "Gagal memenuhi kewajiban (kerja/sekolah/rumah)",
    "Terus menggunakan meski ada masalah sosial/interpersonal",
    "Mengurangi/meninggalkan aktivitas penting karena penggunaan",
    "Menggunakan dalam situasi berbahaya",
    "Terus menggunakan meski tahu ada masalah fisik/psikologis",
    "Toleransi (butuh dosis lebih tinggi)",
    "Withdrawal/Sakau (gejala putus zat)"
]

POLA_PENGGUNAAN = ["Coba-Coba", "Rekreasional", "Situasional", "Habitual", "Kompulsif"]
TINGKAT_KEPARAHAN = ["Ringan", "Sedang", "Berat"]

# =============================================================================
# FUNGSI GENERATE NOMOR SURAT
# =============================================================================
def generate_nomor_surat():
    """Generate nomor surat otomatis"""
    now = datetime.now()
    return f"B/{now.strftime('%m.%d')}/X/KA/PB.06/{now.year}/BNN KALTARA"

# =============================================================================
# FUNGSI ANALISIS MEDIS
# =============================================================================
def analyze_medical_data(data):
    """Analisis data medis berdasarkan DSM-5 dan ASI"""
    dsm5_count = data.get('dsm5_count', 0)
    
    # Tentukan tingkat keparahan berdasarkan DSM-5
    if dsm5_count <= 1:
        severity = "Tidak ada gangguan"
        severity_level = "Tidak Ada"
    elif dsm5_count <= 3:
        severity = "Gangguan Penggunaan Ringan (Mild)"
        severity_level = "Ringan"
    elif dsm5_count <= 5:
        severity = "Gangguan Penggunaan Sedang (Moderate)"
        severity_level = "Sedang"
    else:
        severity = "Gangguan Penggunaan Berat (Severe)"
        severity_level = "Berat"
    
    # Tentukan diagnosis ICD-10
    jenis_utama = data.get('jenis_narkotika_utama', '')
    if 'Sabu' in jenis_utama or 'Ekstasi' in jenis_utama:
        diagnosis_code = "F15"
    elif 'Heroin' in jenis_utama or 'Morfin' in jenis_utama:
        diagnosis_code = "F11"
    elif 'Ganja' in jenis_utama or 'Cannabinoid' in jenis_utama:
        diagnosis_code = "F12"
    elif 'Kokain' in jenis_utama:
        diagnosis_code = "F14"
    else:
        diagnosis_code = "F19"
    
    diagnosis = DIAGNOSIS_ICD10.get(diagnosis_code, "Gangguan Mental dan Perilaku akibat Penggunaan Zat")
    
    return {
        'dsm5_count': dsm5_count,
        'severity': severity,
        'severity_level': severity_level,
        'diagnosis_code': diagnosis_code,
        'diagnosis': diagnosis,
        'pola_penggunaan': data.get('pola_penggunaan', 'Situasional'),
        'durasi_bulan': data.get('durasi_bulan', 0)
    }

# =============================================================================
# FUNGSI ANALISIS HUKUM
# =============================================================================
def analyze_legal_data(data):
    """Analisis data hukum berdasarkan KEP/99"""
    
    # Cek keterlibatan jaringan
    tujuan_kepemilikan = data.get('tujuan_kepemilikan', '')
    metode_pembelian = data.get('metode_pembelian', '')
    
    if tujuan_kepemilikan in ['Dipakai Sendiri', 'Dipakai Bersama-sama'] and \
       'Jaringan Tertentu' not in metode_pembelian:
        keterlibatan_jaringan = "Tidak didapatkan"
    else:
        keterlibatan_jaringan = "Didapatkan"
    
    # Cek riwayat pidana
    riwayat_pidana = data.get('riwayat_pidana_narkotika', False)
    riwayat_penahanan = data.get('riwayat_penahanan', 0)
    
    return {
        'keterlibatan_jaringan': keterlibatan_jaringan,
        'riwayat_pidana': riwayat_pidana,
        'riwayat_penahanan': riwayat_penahanan,
        'barang_bukti': data.get('barang_bukti_jenis', []),
        'tujuan_kepemilikan': tujuan_kepemilikan
    }

# =============================================================================
# FUNGSI REKOMENDASI
# =============================================================================
def generate_recommendation(medical_analysis, legal_analysis, demografi):
    """Generate rekomendasi berdasarkan analisis medis dan hukum"""
    
    severity = medical_analysis['severity_level']
    keterlibatan = legal_analysis['keterlibatan_jaringan']
    dsm5 = medical_analysis['dsm5_count']
    
    # Rule-based recommendation
    if keterlibatan == "Tidak didapatkan" and dsm5 >= 2:
        if severity == "Berat" or dsm5 >= 6:
            rekomendasi = "Rehabilitasi Rawat Inap"
            durasi = "6 (enam) bulan"
            tempat = "RS/Balai Besar Rehabilitasi/Lembaga Rehabilitasi/Institusi Penerima Wajib Lapor Badan Narkotika Nasional"
        else:
            rekomendasi = "Rehabilitasi Rawat Jalan"
            durasi = "3 (tiga) bulan"
            tempat = "Institusi Penerima Wajib Lapor Badan Narkotika Nasional"
            
        tindak_lanjut = "dilanjutkan sesuai ketentuan Perundang-Undangan"
        wajib_lapor = f"melaksanakan WAJIB LAPOR kepada Penyidik {demografi.get('instansi_penyidik', 'Polda/Polres')} sampai selesai proses rehabilitasi"
        
    elif keterlibatan == "Didapatkan" and dsm5 >= 2:
        rekomendasi = "Proses Hukum dengan Rehabilitasi"
        durasi = "sesuai putusan hakim"
        tempat = "Lembaga Pemasyarakatan dengan fasilitas rehabilitasi"
        tindak_lanjut = "dilanjutkan proses hukum dengan mempertimbangkan aspek rehabilitasi"
        wajib_lapor = "menjalani rehabilitasi dalam masa penahanan/pidana"
        
    else:
        rekomendasi = "Proses Hukum"
        durasi = "-"
        tempat = "-"
        tindak_lanjut = "dilanjutkan sesuai ketentuan Perundang-Undangan"
        wajib_lapor = "-"
    
    return {
        'rekomendasi': rekomendasi,
        'durasi': durasi,
        'tempat': tempat,
        'tindak_lanjut': tindak_lanjut,
        'wajib_lapor': wajib_lapor
    }

# =============================================================================
# FUNGSI GENERATE WORD DOCUMENT
# =============================================================================
def generate_word_document(data, medical_analysis, legal_analysis, recommendation):
    """Generate dokumen Word format surat TAT"""
    
    doc = Document()
    
    # Setup margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)
    
    # Header
    header_text = doc.add_paragraph()
    header_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = header_text.add_run("BADAN NARKOTIKA NASIONAL\nPROVINSI KALIMANTAN UTARA\n")
    run.bold = True
    run.font.size = Pt(14)
    
    run2 = header_text.add_run("(NATIONAL NARCOTICS BOARD PROVINCE OF NORTH KALIMANTAN)\n")
    run2.font.size = Pt(11)
    
    run3 = header_text.add_run("Jl. Teuku Umar No. 31, Kota Tarakan, Provinsi Kalimantan Utara\n")
    run3.font.size = Pt(10)
    
    run4 = header_text.add_run("Telepon: (0551) 21XXX | Email: bnnkaltara@bnn.go.id")
    run4.font.size = Pt(9)
    
    # Garis pemisah
    doc.add_paragraph("_" * 80)
    
    # Nomor surat
    p = doc.add_paragraph()
    p.add_run(f"Nomor\t\t: {data['nomor_surat']}\n")
    p.add_run(f"Klasifikasi\t: RAHASIA\n")
    p.add_run(f"Lampiran\t: 1 (satu) berkas BA TAT\n")
    p.add_run(f"Perihal\t\t: Hasil Asesmen Terpadu Tersangka a.n. {data['nama']}")
    
    p_right = doc.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run(f"Tarakan, {data['tanggal_surat']}")
    
    doc.add_paragraph()
    
    # Kepada
    p_kepada = doc.add_paragraph()
    p_kepada.add_run("Kepada\nYth. ").bold = True
    p_kepada.add_run(f"{data['penerima_surat']}\ndi\nTempat")
    
    doc.add_paragraph()
    
    # Rujukan
    doc.add_paragraph("1. Rujukan:").bold = True
    rujukan_list = [
        "Undang-Undang Nomor 35 Tahun 2009 tentang Narkotika;",
        "Peraturan Presiden Nomor 47 Tahun 2019 tentang Perubahan atas Peraturan Presiden Nomor 23 Tahun 2010 tentang Badan Narkotika Nasional;",
        "Peraturan Bersama 7 Instansi Nomor 1 Tahun 2014 tentang Penanganan Pecandu dan Korban Penyalahgunaan Narkotika ke dalam Lembaga Rehabilitasi;",
        "Keputusan Kepala Badan Narkotika Nasional Nomor KEP/99 I/X/KA/PB/06.00/2025/BNN tentang Petunjuk Teknis Pelaksanaan Asesmen Terpadu;",
        f"Surat {data['instansi_pemohon']} Nomor: {data['nomor_surat_pemohon']} tanggal {data['tanggal_surat_pemohon']} perihal Permohonan Asesmen Terpadu."
    ]
    
    for idx, item in enumerate(rujukan_list, start=1):
        doc.add_paragraph(f"{chr(96+idx)}. {item}", style='List Bullet')
    
    doc.add_paragraph()
    
    # Isi
    doc.add_paragraph("2. Sehubungan dengan rujukan tersebut di atas, bersama ini disampaikan bahwa Tim Asesmen Terpadu Tingkat Provinsi Kalimantan Utara telah melakukan Asesmen Terpadu pada:")
    
    # Data tersangka
    table_data = [
        ["Nama", f": {data['nama']}"],
        ["NIK", f": {data['nik']}"],
        ["Tempat/Tgl Lahir", f": {data['tempat_lahir']}, {data['tanggal_lahir']}"],
        ["Jenis Kelamin", f": {data['jenis_kelamin']}"],
        ["Kewarganegaraan", f": {data['kewarganegaraan']}"],
        ["Alamat", f": {data['alamat']}"],
    ]
    
    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Light List'
    
    for idx, (label, value) in enumerate(table_data):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value
    
    doc.add_paragraph()
    
    # Kesimpulan
    doc.add_paragraph("3. Berdasarkan hasil Asesmen Terpadu terhadap tersangka/terdakwa a.n. ").add_run(data['nama']).bold = True
    doc.add_paragraph("Tim Asesmen Terpadu Tingkat Provinsi Kalimantan Utara menyimpulkan:").add_run().bold = True
    
    # Poin a - Kesimpulan Medis
    jenis_narkotika_text = ", ".join(data.get('jenis_narkotika_positif', []))
    
    p_a = doc.add_paragraph()
    p_a.add_run(f"a. Bahwa tersangka/terdakwa a.n. {data['nama']} merupakan penyalahguna narkotika golongan I yaitu {jenis_narkotika_text} untuk diri sendiri dengan pola pemakaian {medical_analysis['pola_penggunaan']} kategori {medical_analysis['severity_level']}, didiagnosis {medical_analysis['diagnosis']} ({medical_analysis['diagnosis_code']}).")
    
    # Poin b - Kesimpulan Hukum
    p_b = doc.add_paragraph()
    p_b.add_run(f"b. Bahwa tersangka/terdakwa a.n. {data['nama']} {legal_analysis['keterlibatan_jaringan']} indikasi keterlibatan dalam jaringan peredaran gelap narkotika.")
    
    doc.add_paragraph()
    
    # Rekomendasi
    doc.add_paragraph("4. Tim Asesmen Terpadu Tingkat Provinsi Kalimantan Utara memberikan rekomendasi terhadap ").add_run(f"tersangka/terdakwa a.n. {data['nama']}").bold = True
    doc.add_paragraph(" sebagai berikut:")
    
    p_rec_a = doc.add_paragraph()
    if recommendation['rekomendasi'] != "Proses Hukum":
        p_rec_a.add_run(f"a. Terhadap tersangka/terdakwa a.n. {data['nama']} agar dilakukan perawatan dan pemulihan dengan {recommendation['rekomendasi']} sebanyak {recommendation['durasi']} di {recommendation['tempat']} dan {recommendation['wajib_lapor']}.")
    else:
        p_rec_a.add_run(f"a. Terhadap tersangka/terdakwa a.n. {data['nama']} agar {recommendation['tindak_lanjut']}.")
    
    p_rec_b = doc.add_paragraph()
    p_rec_b.add_run(f"b. Terhadap perkara tersangka/terdakwa a.n. {data['nama']} {recommendation['tindak_lanjut']}.")
    
    doc.add_paragraph()
    
    # Penutup
    doc.add_paragraph("5. Demikian untuk menjadi periksa.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_ttd.add_run("Ditandatangani Secara Elektronik Oleh:\n\n")
    p_ttd.add_run(f"{data['jabatan_penandatangan']}\n\n\n\n")
    p_ttd.add_run(f"{data['nama_penandatangan']}\n")
    p_ttd.add_run(f"NIP. {data['nip_penandatangan']}")
    
    # Tembusan
    doc.add_paragraph("\n\nTembusan:")
    tembusan_list = [
        "Kepala BNN Provinsi Kalimantan Utara;",
        "Sekretaris BNN Provinsi Kalimantan Utara;",
        "Kepala Seksi Rehabilitasi BNN Provinsi Kalimantan Utara;",
        f"{data['instansi_pemohon']}."
    ]
    
    for idx, item in enumerate(tembusan_list, start=1):
        doc.add_paragraph(f"{idx}. {item}", style='List Number')
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================================================================
# FUNGSI GENERATE PDF
# =============================================================================
def generate_pdf_document(data, medical_analysis, legal_analysis, recommendation):
    """Generate dokumen PDF format surat TAT"""
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           topMargin=1*cm, bottomMargin=1*cm,
                           leftMargin=1.5*cm, rightMargin=1.5*cm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_center = ParagraphStyle(
        'CustomCenter',
        parent=styles['Normal'],
        alignment=TA_CENTER,
        fontSize=12,
        spaceAfter=6
    )
    
    style_header = ParagraphStyle(
        'CustomHeader',
        parent=styles['Normal'],
        alignment=TA_CENTER,
        fontSize=14,
        fontName='Helvetica-Bold',
        spaceAfter=6
    )
    
    style_body = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14,
        spaceAfter=8
    )
    
    # Header
    elements.append(Paragraph("<b>BADAN NARKOTIKA NASIONAL<br/>PROVINSI KALIMANTAN UTARA</b>", style_header))
    elements.append(Paragraph("(NATIONAL NARCOTICS BOARD PROVINCE OF NORTH KALIMANTAN)", style_center))
    elements.append(Paragraph("Jl. Teuku Umar No. 31, Kota Tarakan, Provinsi Kalimantan Utara", style_center))
    elements.append(Paragraph("Telepon: (0551) 21XXX | Email: bnnkaltara@bnn.go.id", style_center))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph("_" * 100, style_center))
    elements.append(Spacer(1, 0.5*cm))
    
    # Nomor surat
    nomor_text = f"""
    <b>Nomor</b> : {data['nomor_surat']}<br/>
    <b>Klasifikasi</b> : RAHASIA<br/>
    <b>Lampiran</b> : 1 (satu) berkas BA TAT<br/>
    <b>Perihal</b> : Hasil Asesmen Terpadu Tersangka a.n. {data['nama']}
    """
    elements.append(Paragraph(nomor_text, style_body))
    elements.append(Spacer(1, 0.3*cm))
    
    tanggal_style = ParagraphStyle('TanggalRight', parent=styles['Normal'], alignment=2, fontSize=11)
    elements.append(Paragraph(f"Tarakan, {data['tanggal_surat']}", tanggal_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Kepada
    kepada_text = f"<b>Kepada</b><br/>Yth. {data['penerima_surat']}<br/>di<br/>Tempat"
    elements.append(Paragraph(kepada_text, style_body))
    elements.append(Spacer(1, 0.5*cm))
    
    # Isi surat
    elements.append(Paragraph("<b>1. Rujukan:</b>", style_body))
    
    rujukan = [
        "Undang-Undang Nomor 35 Tahun 2009 tentang Narkotika;",
        "Peraturan Bersama 7 Instansi Nomor 1 Tahun 2014;",
        "KEP/99 I/X/KA/PB/06.00/2025/BNN;",
        f"Surat {data['instansi_pemohon']} Nomor: {data['nomor_surat_pemohon']}."
    ]
    
    for idx, r in enumerate(rujukan, 1):
        elements.append(Paragraph(f"{chr(96+idx)}. {r}", style_body))
    
    elements.append(Spacer(1, 0.3*cm))
    
    # Data tersangka
    elements.append(Paragraph(f"<b>2.</b> Tim Asesmen Terpadu telah melakukan asesmen terhadap:", style_body))
    
    data_table = [
        ["Nama", f": {data['nama']}"],
        ["NIK", f": {data['nik']}"],
        ["Tempat/Tgl Lahir", f": {data['tempat_lahir']}, {data['tanggal_lahir']}"],
        ["Jenis Kelamin", f": {data['jenis_kelamin']}"],
        ["Alamat", f": {data['alamat']}"]
    ]
    
    t = Table(data_table, colWidths=[4*cm, 12*cm])
    t.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.5*cm))
    
    # Kesimpulan
    jenis_narkotika = ", ".join(data.get('jenis_narkotika_positif', []))
    
    elements.append(Paragraph(f"<b>3. Kesimpulan:</b>", style_body))
    elements.append(Paragraph(
        f"a. Tersangka/terdakwa a.n. <b>{data['nama']}</b> merupakan penyalahguna narkotika "
        f"({jenis_narkotika}) untuk diri sendiri dengan pola {medical_analysis['pola_penggunaan']} "
        f"kategori {medical_analysis['severity_level']}, didiagnosis {medical_analysis['diagnosis']} "
        f"({medical_analysis['diagnosis_code']}).",
        style_body
    ))
    
    elements.append(Paragraph(
        f"b. Tersangka/terdakwa <b>{legal_analysis['keterlibatan_jaringan']}</b> indikasi "
        f"keterlibatan dalam jaringan peredaran gelap narkotika.",
        style_body
    ))
    
    elements.append(Spacer(1, 0.3*cm))
    
    # Rekomendasi
    elements.append(Paragraph(f"<b>4. Rekomendasi:</b>", style_body))
    
    if recommendation['rekomendasi'] != "Proses Hukum":
        elements.append(Paragraph(
            f"a. Dilakukan {recommendation['rekomendasi']} sebanyak {recommendation['durasi']} "
            f"di {recommendation['tempat']} dan {recommendation['wajib_lapor']}.",
            style_body
        ))
    else:
        elements.append(Paragraph(
            f"a. {recommendation['tindak_lanjut']}.",
            style_body
        ))
    
    elements.append(Paragraph(
        f"b. Terhadap perkara {recommendation['tindak_lanjut']}.",
        style_body
    ))
    
    elements.append(Spacer(1, 0.5*cm))
    elements.append(Paragraph("<b>5. Demikian untuk menjadi periksa.</b>", style_body))
    elements.append(Spacer(1, 1*cm))
    
    # Tanda tangan
    ttd_text = f"""
    Ditandatangani Secara Elektronik Oleh:<br/><br/>
    <b>{data['jabatan_penandatangan']}</b><br/><br/><br/><br/>
    <b>{data['nama_penandatangan']}</b><br/>
    NIP. {data['nip_penandatangan']}
    """
    elements.append(Paragraph(ttd_text, style_body))
    
    elements.append(Spacer(1, 0.5*cm))
    
    # Tembusan
    elements.append(Paragraph("<b>Tembusan:</b>", style_body))
    tembusan = [
        "Kepala BNN Provinsi Kalimantan Utara;",
        "Sekretaris BNN Provinsi Kalimantan Utara;",
        "Kepala Seksi Rehabilitasi BNN Provinsi Kalimantan Utara;",
        f"{data['instansi_pemohon']}."
    ]
    
    for idx, t in enumerate(tembusan, 1):
        elements.append(Paragraph(f"{idx}. {t}", style_body))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

# =============================================================================
# MAIN APPLICATION
# =============================================================================
def main():
    st.markdown('<h1 class="main-header">⚖️ SISTEM ASESMEN TERPADU (TAT)<br/>BNN PROVINSI KALIMANTAN UTARA</h1>', 
                unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
        <strong>📋 Berdasarkan:</strong><br/>
        • UU No. 35 Tahun 2009 tentang Narkotika<br/>
        • KEP/99 I/X/KA/PB/06.00/2025/BNN tentang Petunjuk Teknis Pelaksanaan Asesmen Terpadu<br/>
        • Instrumen: ASI, ASAM, DSM-5, ICD-10, PPDGJ III
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("📌 Informasi Sistem")
        st.info("""
        **BNN Provinsi Kalimantan Utara**
        
        Jl. Teuku Umar No. 31
        Kota Tarakan
        Provinsi Kalimantan Utara
        
        📞 (0551) 21XXX
        📧 bnnkaltara@bnn.go.id
        """)
        
        st.markdown("---")
        st.caption("Versi 2.0 - Desember 2025")
    
    # Tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "📝 I. DEMOGRAFI & HUKUM", 
        "🏥 II. ASESMEN MEDIS",
        "📄 III. HASIL & SURAT TAT",
        "📚 IV. PANDUAN"
    ])
    
    # =============================================================================
    # TAB 1: DEMOGRAFI & ASESMEN HUKUM
    # =============================================================================
    with tab1:
        st.header("📋 I. DATA DEMOGRAFI DAN ASESMEN HUKUM")
        
        st.markdown("---")
        st.subheader("A. IDENTITAS TERSANGKA/TERDAKWA")
        
        col1, col2 = st.columns(2)
        
        with col1:
            nama = st.text_input("Nama Lengkap *", placeholder="Contoh: AHMAD YANI")
            nik = st.text_input("NIK *", placeholder="6471XXXXXXXXXXXX", max_chars=16)
            tempat_lahir = st.text_input("Tempat Lahir *", placeholder="Tarakan")
            tanggal_lahir = st.date_input("Tanggal Lahir *", 
                                         min_value=datetime(1950, 1, 1),
                                         max_value=datetime.now())
            jenis_kelamin = st.selectbox("Jenis Kelamin *", ["Laki-laki", "Perempuan"])
            kewarganegaraan = st.text_input("Kewarganegaraan *", value="Indonesia")
        
        with col2:
            alamat = st.text_area("Alamat Lengkap *", 
                                 placeholder="Jl. ..., RT/RW, Kelurahan, Kecamatan, Kota/Kab, Provinsi",
                                 height=100)
            no_hp = st.text_input("Nomor HP", placeholder="08XXXXXXXXXX")
            no_rekening = st.text_input("Nomor Rekening", placeholder="Bank ... No. ...")
            status_kawin = st.selectbox("Status Perkawinan", 
                                       ["Belum Kawin", "Kawin", "Cerai Hidup", "Cerai Mati"])
        
        col3, col4 = st.columns(2)
        
        with col3:
            pendidikan = st.selectbox("Pendidikan Terakhir",
                                     ["Tidak Sekolah", "SD", "SMP", "SMA/SMK", 
                                      "D3", "S1", "S2", "S3"])
            pekerjaan = st.text_input("Pekerjaan Saat Ini", placeholder="Contoh: Karyawan Swasta")
        
        with col4:
            penghasilan = st.number_input("Rata-rata Penghasilan/Bulan (Rp)", 
                                         min_value=0, value=0, step=100000,
                                         format="%d")
        
        catatan_demografi = st.text_area("Catatan Tambahan Demografi", 
                                        placeholder="Informasi tambahan yang relevan...",
                                        height=80)
        
        st.markdown("---")
        st.subheader("B. KRONOLOGI KEJADIAN")
        
        kronologi = st.text_area(
            "Uraikan kronologi penangkapan/penyerahan diri *",
            placeholder="Contoh: Pada hari Senin tanggal ... tersangka ditangkap di ... "
                       "saat melakukan ... Barang bukti yang ditemukan berupa ...",
            height=150
        )
        
        st.markdown("---")
        st.subheader("C. PENGGUNAAN NARKOTIKA")
        
        st.markdown("**1. Jenis Narkotika yang Digunakan**")
        
        jenis_narkotika_digunakan = st.multiselect(
            "Pilih jenis narkotika yang pernah/sedang digunakan *",
            JENIS_NARKOTIKA_LIST
        )
        
        if "Lainnya" in jenis_narkotika_digunakan:
            jenis_lainnya = st.text_input("Sebutkan jenis lainnya:")
        
        st.markdown("**2. Hasil Pemeriksaan Urine/Laboratorium**")
        
        col_urine1, col_urine2 = st.columns(2)
        
        with col_urine1:
            hasil_urine = st.radio("Hasil Tes Urine *", ["Positif", "Negatif"])
        
        with col_urine2:
            if hasil_urine == "Positif":
                jenis_positif = st.multiselect(
                    "Jenis Narkotika yang Positif *",
                    JENIS_NARKOTIKA_LIST
                )
            else:
                jenis_positif = []
        
        st.markdown("---")
        st.subheader("D. STATUS HUKUM")
        
        st.markdown("**1. Riwayat Tindak Pidana**")
        
        col_pidana1, col_pidana2, col_pidana3 = st.columns(3)
        
        with col_pidana1:
            riwayat_narkotika = st.number_input("Narkotika (kali)", min_value=0, value=0)
            riwayat_psikotropika = st.number_input("Psikotropika (kali)", min_value=0, value=0)
            riwayat_pencurian = st.number_input("Pencurian (kali)", min_value=0, value=0)
        
        with col_pidana2:
            riwayat_perampokan = st.number_input("Perampokan (kali)", min_value=0, value=0)
            riwayat_pembunuhan = st.number_input("Pembunuhan (kali)", min_value=0, value=0)
            riwayat_pemerkosaan = st.number_input("Pemerkosaan (kali)", min_value=0, value=0)
        
        with col_pidana3:
            riwayat_lainnya_pidana = st.text_input("Tindak Pidana Lainnya")
            jumlah_lainnya = st.number_input("Jumlah (kali)", min_value=0, value=0)
        
        st.markdown("**2. Riwayat Penahanan**")
        
        col_tahan1, col_tahan2 = st.columns(2)
        
        with col_tahan1:
            jumlah_penahanan = st.number_input("Jumlah Penahanan (kali)", min_value=0, value=0)
            
            if jumlah_penahanan > 0:
                tempat_penahanan = st.text_input("Tempat Penahanan Terakhir", 
                                                placeholder="Contoh: Polres Tarakan")
                tanggal_penahanan = st.date_input("Tanggal Penahanan")
        
        with col_tahan2:
            if jumlah_penahanan > 0:
                lama_penahanan = st.number_input("Lama Penahanan (hari)", min_value=0, value=0)
                
                status_penahanan = st.selectbox("Status Akhir Penahanan",
                                               ["Penangguhan Penahanan",
                                                "Bebas Demi Hukum",
                                                "Proses Hukum Lanjut"])
        
        st.markdown("**3. Riwayat Persidangan**")
        
        pernah_sidang = st.checkbox("Pernah menjalani persidangan kasus narkotika?")
        
        if pernah_sidang:
            col_sidang1, col_sidang2 = st.columns(2)
            
            with col_sidang1:
                tindak_pidana_sidang = st.text_input("Tindak Pidana yang Disidangkan")
                vonis_tahun = st.number_input("Vonis Hakim (tahun)", min_value=0.0, value=0.0, step=0.5)
            
            with col_sidang2:
                tempat_vonis = st.text_input("Ditempatkan di", placeholder="Rutan/Lapas ...")
        
        st.markdown("---")
        st.subheader("E. KETERLIBATAN DALAM JARINGAN")
        
        st.markdown("**1. Narkotika yang Dimiliki Saat Penangkapan**")
        
        barang_bukti_jenis = st.multiselect(
            "Jenis narkotika yang menjadi barang bukti *",
            JENIS_NARKOTIKA_LIST
        )
        
        barang_bukti_detail = {}
        for jenis in barang_bukti_jenis:
            col_bb1, col_bb2 = st.columns(2)
            with col_bb1:
                jumlah = st.number_input(f"Jumlah {jenis} (gram)", 
                                        min_value=0.0, value=0.0, step=0.1,
                                        key=f"bb_{jenis}")
            with col_bb2:
                satuan = st.selectbox(f"Satuan {jenis}", 
                                     ["gram", "butir", "paket", "lainnya"],
                                     key=f"satuan_{jenis}")
            
            barang_bukti_detail[jenis] = {"jumlah": jumlah, "satuan": satuan}
        
        st.markdown("**2. Tujuan Kepemilikan Narkotika**")
        
        tujuan_kepemilikan = st.radio(
            "Narkotika yang dimiliki untuk *",
            ["Dipakai Sendiri",
             "Dipakai Bersama-sama",
             "Titipan Orang",
             "Akan Dijual",
             "Lainnya"]
        )
        
        if tujuan_kepemilikan == "Lainnya":
            tujuan_lainnya = st.text_input("Sebutkan:")
        
        st.markdown("**3. Metode Pembelian Narkotika**")
        
        col_metode1, col_metode2 = st.columns(2)
        
        with col_metode1:
            metode_pembelian = st.selectbox(
                "Cara mendapatkan narkotika *",
                ["Beli Langsung di Tempat",
                 "Dari Teman",
                 "Dari Jaringan Tertentu",
                 "Aplikasi/Sosial Media",
                 "Lainnya"]
            )
            
            if metode_pembelian == "Beli Langsung di Tempat":
                lokasi_beli = st.text_input("Lokasi pembelian:")
            
            metode_pembayaran = st.selectbox(
                "Metode Pembayaran",
                ["Cash", "Transfer Bank", "Uang Elektronik", "Lainnya"]
            )
        
        with col_metode2:
            untuk_siapa = st.selectbox(
                "Narkotika dibeli untuk",
                ["Diri Sendiri", "Teman", "Orang Lain"]
            )
            
            frekuensi_beli = st.number_input("Frekuensi Pembelian (kali)", 
                                            min_value=0, value=1)
            
            harga_beli = st.number_input("Harga Pembelian Terakhir (Rp)",
                                        min_value=0, value=0, step=50000)
        
        if metode_pembayaran == "Transfer Bank":
            bukti_transfer = st.radio("Bukti Transfer", ["Ada", "Tidak Ada"])
        
        st.markdown("**4. Database Intelijen**")
        
        cek_database = st.checkbox("Sudah dilakukan pengecekan database intelijen?")
        
        if cek_database:
            hasil_database = st.text_area(
                "Hasil Pengecekan Database Intelijen",
                placeholder="Uraikan temuan dari database intelijen (jika ada)...",
                height=100
            )
        
        st.markdown("**5. Fakta-Fakta Hukum & Kesimpulan Hukum**")
        
        fakta_hukum = st.text_area(
            "Fakta-Fakta Hukum *",
            placeholder="Uraikan fakta-fakta hukum yang ditemukan dalam penyidikan...",
            height=120
        )
        
        kesimpulan_hukum = st.text_area(
            "Kesimpulan Asesmen Hukum *",
            placeholder="Kesimpulan keterlibatan dalam jaringan peredaran gelap narkotika...",
            height=100
        )
    
    # =============================================================================
    # TAB 2: ASESMEN MEDIS
    # =============================================================================
    with tab2:
        st.header("🏥 II. ASESMEN MEDIS")
        
        st.markdown("""
        <div class="info-box">
        <strong>📋 Instrumen Asesmen:</strong><br/>
        • DSM-5 (Diagnostic and Statistical Manual of Mental Disorders)<br/>
        • ICD-10 / PPDGJ III (International Classification of Diseases)<br/>
        • ASAM (American Society of Addiction Medicine)<br/>
        • ASI (Addiction Severity Index)
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.subheader("A. KRITERIA DSM-5 (Gangguan Penggunaan Zat)")
        
        st.markdown("""
        **Petunjuk:** Berikan tanda centang pada kriteria yang **TERPENUHI** berdasarkan 
        wawancara dan observasi klinis.
        """)
        
        dsm5_checked = []
        
        for i, criteria in enumerate(DSM5_CRITERIA, 1):
            if st.checkbox(f"{i}. {criteria}", key=f"dsm5_med_{i}"):
                dsm5_checked.append(criteria)
        
        dsm5_count = len(dsm5_checked)
        
        # Interpretasi DSM-5
        col_dsm1, col_dsm2, col_dsm3 = st.columns(3)
        
        with col_dsm1:
            st.metric("Kriteria Terpenuhi", f"{dsm5_count}/11")
        
        with col_dsm2:
            if dsm5_count <= 1:
                st.info("✓ Tidak ada gangguan")
                severity_auto = "Tidak Ada"
            elif dsm5_count <= 3:
                st.warning("⚠ Gangguan RINGAN")
                severity_auto = "Ringan"
            elif dsm5_count <= 5:
                st.warning("⚠ Gangguan SEDANG")
                severity_auto = "Sedang"
            else:
                st.error("⚠ Gangguan BERAT")
                severity_auto = "Berat"
        
        with col_dsm3:
            st.info(f"**Kategori:** {severity_auto}")
        
        st.markdown("---")
        st.subheader("B. DIAGNOSIS ICD-10 / PPDGJ III")
        
        jenis_utama_medis = st.selectbox(
            "Jenis Narkotika Utama yang Digunakan *",
            JENIS_NARKOTIKA_LIST,
            help="Pilih jenis narkotika yang paling dominan/sering digunakan"
        )
        
        # Auto-suggest diagnosis berdasarkan jenis narkotika
        if 'Sabu' in jenis_utama_medis or 'Ekstasi' in jenis_utama_medis:
            diagnosis_suggest = "F15"
        elif 'Heroin' in jenis_utama_medis or 'Morfin' in jenis_utama_medis:
            diagnosis_suggest = "F11"
        elif 'Ganja' in jenis_utama_medis or 'Cannabinoid' in jenis_utama_medis:
            diagnosis_suggest = "F12"
        elif 'Kokain' in jenis_utama_medis:
            diagnosis_suggest = "F14"
        else:
            diagnosis_suggest = "F19"
        
        diagnosis_code = st.selectbox(
            "Kode Diagnosis ICD-10 *",
            list(DIAGNOSIS_ICD10.keys()),
            index=list(DIAGNOSIS_ICD10.keys()).index(diagnosis_suggest)
        )
        
        st.info(f"**Diagnosis:** {DIAGNOSIS_ICD10[diagnosis_code]}")
        
        st.markdown("---")
        st.subheader("C. POLA PENGGUNAAN NARKOTIKA")
        
        col_pola1, col_pola2 = st.columns(2)
        
        with col_pola1:
            pola_penggunaan = st.selectbox(
                "Pola Penggunaan *",
                POLA_PENGGUNAAN,
                help="Coba-Coba: <5x | Rekreasional: Sesekali di pesta | "
                     "Situasional: Situasi tertentu | Habitual: Rutin | Kompulsif: Tidak terkontrol"
            )
            
            durasi_penggunaan = st.number_input(
                "Durasi Penggunaan (bulan) *",
                min_value=0, value=6, max_value=600
            )
        
        with col_pola2:
            frekuensi_penggunaan = st.selectbox(
                "Frekuensi Penggunaan",
                ["Setiap hari", "3-6x per minggu", "1-2x per minggu", 
                 "1-3x per bulan", "Jarang (< 1x per bulan)"]
            )
            
            cara_penggunaan = st.multiselect(
                "Cara Penggunaan",
                ["Dihisap", "Dihirup", "Diminum", "Disuntik", "Lainnya"]
            )
        
        st.markdown("---")
        st.subheader("D. ASAM 6 DIMENSI (Simplified)")
        
        st.markdown("**1. Dimensi Intoksikasi & Withdrawal**")
        
        col_asam1, col_asam2 = st.columns(2)
        
        with col_asam1:
            ada_withdrawal = st.checkbox("Ada gejala putus zat (sakau/withdrawal)?")
            
            if ada_withdrawal:
                tingkat_withdrawal = st.select_slider(
                    "Tingkat Keparahan Withdrawal",
                    options=["Ringan", "Sedang", "Berat"]
                )
        
        with col_asam2:
            ada_intoksikasi = st.checkbox("Pernah mengalami intoksikasi akut/overdosis?")
            
            if ada_intoksikasi:
                frekuensi_intoksikasi = st.number_input(
                    "Berapa kali?", min_value=1, value=1
                )
        
        st.markdown("**2. Dimensi Kondisi Biomedis**")
        
        ada_penyakit = st.checkbox("Ada penyakit fisik/medis yang menyertai?")
        
        if ada_penyakit:
            jenis_penyakit = st.multiselect(
                "Jenis Penyakit/Kondisi Medis",
                ["HIV/AIDS", "Hepatitis", "TBC", "Penyakit Jantung", 
                 "Diabetes", "Penyakit Kulit", "Lainnya"]
            )
            
            if "Lainnya" in jenis_penyakit:
                penyakit_lainnya = st.text_input("Sebutkan:")
        
        st.markdown("**3. Dimensi Kondisi Emosional/Psikiatrik**")
        
        ada_gangguan_jiwa = st.checkbox("Ada gangguan kesehatan mental/psikiatrik?")
        
        if ada_gangguan_jiwa:
            jenis_gangguan = st.multiselect(
                "Jenis Gangguan Mental",
                ["Depresi", "Anxietas/Kecemasan", "Gangguan Bipolar", 
                 "Skizofrenia", "PTSD", "Gangguan Kepribadian", "Lainnya"]
            )
            
            tingkat_gangguan_jiwa = st.select_slider(
                "Tingkat Keparahan",
                options=["Ringan", "Sedang", "Berat"]
            )
        
        st.markdown("**4. Dimensi Kesiapan Berubah**")
        
        motivasi_rehabilitasi = st.select_slider(
            "Motivasi untuk Rehabilitasi *",
            options=["Sangat Rendah", "Rendah", "Sedang", "Tinggi", "Sangat Tinggi"]
        )
        
        insight_masalah = st.radio(
            "Kesadaran terhadap Masalah Ketergantungan",
            ["Tidak sadar ada masalah (denial)",
             "Mulai menyadari tapi belum siap berubah",
             "Sadar dan siap untuk berubah",
             "Aktif mencari bantuan"]
        )
        
        st.markdown("**5. Dimensi Potensi Relapse**")
        
        col_relapse1, col_relapse2 = st.columns(2)
        
        with col_relapse1:
            riwayat_rehabilitasi = st.number_input(
                "Riwayat Rehabilitasi Sebelumnya (kali)",
                min_value=0, value=0
            )
            
            if riwayat_rehabilitasi > 0:
                hasil_rehabilitasi = st.selectbox(
                    "Hasil Rehabilitasi Terakhir",
                    ["Selesai tapi kambuh (relapse)",
                     "Drop out (tidak selesai)",
                     "Masih dalam proses"]
                )
        
        with col_relapse2:
            trigger_utama = st.multiselect(
                "Pemicu Utama Penggunaan (Trigger)",
                ["Stress/Tekanan", "Lingkungan Pergaulan", "Masalah Keluarga",
                 "Masalah Ekonomi", "Teman Pengguna", "Ketersediaan Narkotika",
                 "Lainnya"]
            )
        
        st.markdown("**6. Dimensi Lingkungan Pemulihan**")
        
        col_ling1, col_ling2 = st.columns(2)
        
        with col_ling1:
            dukungan_keluarga = st.select_slider(
                "Dukungan Keluarga",
                options=["Sangat Tidak Mendukung", "Tidak Mendukung", 
                        "Netral", "Mendukung", "Sangat Mendukung"]
            )
            
            kondisi_rumah = st.selectbox(
                "Kondisi Lingkungan Rumah",
                ["Kondusif untuk pemulihan",
                 "Cukup kondusif",
                 "Tidak kondusif (ada pengguna lain)",
                 "Sangat tidak kondusif (lingkungan peredaran)"]
            )
        
        with col_ling2:
            status_pekerjaan = st.selectbox(
                "Status Pekerjaan/Pendidikan",
                ["Bekerja/Bersekolah aktif",
                 "Tidak bekerja/sekolah tapi produktif",
                 "Tidak bekerja/sekolah tidak produktif",
                 "Kehilangan pekerjaan/DO karena narkotika"]
            )
            
            kemampuan_ekonomi = st.selectbox(
                "Kemampuan Ekonomi untuk Rehabilitasi",
                ["Mampu mandiri",
                 "Mampu dengan bantuan keluarga",
                 "Tidak mampu (butuh bantuan pemerintah)"]
            )
        
        st.markdown("---")
        st.subheader("E. CATATAN KLINIS & KESIMPULAN MEDIS")
        
        catatan_klinis = st.text_area(
            "Catatan Klinis Tambahan",
            placeholder="Observasi perilaku, kondisi fisik saat asesmen, hasil pemeriksaan lain, dll...",
            height=120
        )
        
        kesimpulan_medis = st.text_area(
            "Kesimpulan Asesmen Medis *",
            placeholder="Ringkasan kondisi medis, tingkat kecanduan, diagnosis, dan rekomendasi jenis rehabilitasi...",
            height=150
        )
    
    # =============================================================================
    # TAB 3: HASIL & SURAT TAT
    # =============================================================================
    with tab3:
        st.header("📄 III. HASIL ASESMEN & SURAT TAT")
        
        st.markdown("---")
        st.subheader("A. INFORMASI SURAT")
        
        col_surat1, col_surat2 = st.columns(2)
        
        with col_surat1:
            nomor_surat = st.text_input(
                "Nomor Surat *",
                value=generate_nomor_surat(),
                help="Format otomatis, dapat diubah manual"
            )
            
            tanggal_surat = st.date_input(
                "Tanggal Surat *",
                value=datetime.now()
            )
            
            tanggal_pelaksanaan = st.date_input(
                "Tanggal Pelaksanaan Asesmen *",
                value=datetime.now()
            )
        
        with col_surat2:
            penerima_surat = st.text_input(
                "Penerima Surat (Kepada Yth.) *",
                value="Direktur Reserse Narkoba Polda Kalimantan Utara",
                help="Contoh: Direktur Reserse Narkoba Polda Kaltara / Kapolres Tarakan"
            )
            
            instansi_pemohon = st.text_input(
                "Instansi Pemohon *",
                value="Direktorat Reserse Narkoba Polda Kalimantan Utara"
            )
            
            nomor_surat_pemohon = st.text_input(
                "Nomor Surat Pemohon *",
                placeholder="B/XXX/..."
            )
            
            tanggal_surat_pemohon = st.date_input(
                "Tanggal Surat Pemohon *"
            )
        
        st.markdown("**Penandatangan Surat**")
        
        col_ttd1, col_ttd2, col_ttd3 = st.columns(3)
        
        with col_ttd1:
            jabatan_ttd = st.text_input(
                "Jabatan Penandatangan *",
                value="Kepala Seksi Rehabilitasi BNN Provinsi Kalimantan Utara"
            )
        
        with col_ttd2:
            nama_ttd = st.text_input(
                "Nama Penandatangan *",
                placeholder="Nama Lengkap"
            )
        
        with col_ttd3:
            nip_ttd = st.text_input(
                "NIP*",
                placeholder="19XXXXXX XXXXXX X XXX"
            )

        st.markdown("---")
    
    # Tombol Generate
    if st.button("🔍 PROSES ASESMEN & GENERATE SURAT", type="primary", use_container_width=True):
        
        # Validasi input wajib
        errors = []
        
        if not nama:
            errors.append("Nama lengkap harus diisi")
        if not nik or len(nik) != 16:
            errors.append("NIK harus 16 digit")
        if not alamat:
            errors.append("Alamat harus diisi")
        if not kronologi:
            errors.append("Kronologi kejadian harus diisi")
        if not jenis_narkotika_digunakan:
            errors.append("Jenis narkotika yang digunakan harus dipilih")
        if hasil_urine == "Positif" and not jenis_positif:
            errors.append("Jenis narkotika yang positif harus dipilih")
        if not fakta_hukum:
            errors.append("Fakta-fakta hukum harus diisi")
        if not kesimpulan_hukum:
            errors.append("Kesimpulan hukum harus diisi")
        if not kesimpulan_medis:
            errors.append("Kesimpulan medis harus diisi")
        if not nomor_surat_pemohon:
            errors.append("Nomor surat pemohon harus diisi")
        if not nama_ttd or not nip_ttd:
            errors.append("Data penandatangan harus lengkap")
        
        if errors:
            st.error("⚠️ **Lengkapi data berikut:**")
            for error in errors:
                st.markdown(f"- {error}")
        else:
            with st.spinner("🔄 Memproses asesmen dan membuat surat..."):
                
                # Kompilasi data
                data_lengkap = {
                    # Identitas
                    'nama': nama,
                    'nik': nik,
                    'tempat_lahir': tempat_lahir,
                    'tanggal_lahir': tanggal_lahir.strftime("%d-%m-%Y"),
                    'jenis_kelamin': jenis_kelamin,
                    'kewarganegaraan': kewarganegaraan,
                    'alamat': alamat,
                    'no_hp': no_hp,
                    'no_rekening': no_rekening,
                    'status_kawin': status_kawin,
                    'pendidikan': pendidikan,
                    'pekerjaan': pekerjaan,
                    'penghasilan': penghasilan,
                    
                    # Hukum
                    'kronologi': kronologi,
                    'jenis_narkotika_digunakan': jenis_narkotika_digunakan,
                    'hasil_urine': hasil_urine,
                    'jenis_narkotika_positif': jenis_positif,
                    'riwayat_pidana_narkotika': riwayat_narkotika > 0,
                    'riwayat_penahanan': jumlah_penahanan,
                    'barang_bukti_jenis': barang_bukti_jenis,
                    'barang_bukti_detail': barang_bukti_detail,
                    'tujuan_kepemilikan': tujuan_kepemilikan,
                    'metode_pembelian': metode_pembelian,
                    'fakta_hukum': fakta_hukum,
                    'kesimpulan_hukum': kesimpulan_hukum,
                    
                    # Medis
                    'dsm5_count': dsm5_count,
                    'jenis_narkotika_utama': jenis_utama_medis,
                    'diagnosis_code': diagnosis_code,
                    'pola_penggunaan': pola_penggunaan,
                    'durasi_bulan': durasi_penggunaan,
                    'kesimpulan_medis': kesimpulan_medis,
                    
                    # Surat
                    'nomor_surat': nomor_surat,
                    'tanggal_surat': tanggal_surat.strftime("%d %B %Y"),
                    'tanggal_pelaksanaan': tanggal_pelaksanaan.strftime("%d %B %Y"),
                    'penerima_surat': penerima_surat,
                    'instansi_pemohon': instansi_pemohon,
                    'nomor_surat_pemohon': nomor_surat_pemohon,
                    'tanggal_surat_pemohon': tanggal_surat_pemohon.strftime("%d %B %Y"),
                    'jabatan_penandatangan': jabatan_ttd,
                    'nama_penandatangan': nama_ttd,
                    'nip_penandatangan': nip_ttd,
                    'instansi_penyidik': instansi_pemohon
                }
                
                # Analisis
                medical_analysis = analyze_medical_data(data_lengkap)
                legal_analysis = analyze_legal_data(data_lengkap)
                recommendation = generate_recommendation(medical_analysis, legal_analysis, data_lengkap)
                
                # Simpan ke session state
                st.session_state['hasil_asesmen'] = {
                    'data': data_lengkap,
                    'medical': medical_analysis,
                    'legal': legal_analysis,
                    'recommendation': recommendation,
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                st.success("✅ **Asesmen berhasil diproses!**")
                st.balloons()
    
    # Tampilkan hasil jika sudah ada
    if 'hasil_asesmen' in st.session_state:
        hasil = st.session_state['hasil_asesmen']
        
        st.markdown("---")
        st.subheader("B. RINGKASAN HASIL ASESMEN")
        
        col_hasil1, col_hasil2, col_hasil3 = st.columns(3)
        
        with col_hasil1:
            st.markdown("""
            <div class="info-box">
            <strong>📊 ASESMEN MEDIS</strong><br/>
            • DSM-5: {}/11 kriteria<br/>
            • Tingkat: <strong>{}</strong><br/>
            • Diagnosis: {}<br/>
            • Pola: {}
            </div>
            """.format(
                hasil['medical']['dsm5_count'],
                hasil['medical']['severity_level'],
                hasil['medical']['diagnosis_code'],
                hasil['medical']['pola_penggunaan']
            ), unsafe_allow_html=True)
        
        with col_hasil2:
            st.markdown("""
            <div class="warning-box">
            <strong>⚖️ ASESMEN HUKUM</strong><br/>
            • Keterlibatan Jaringan:<br/>
              <strong>{}</strong><br/>
            • Barang Bukti:<br/>
              {} jenis<br/>
            • Riwayat Pidana: {}
            </div>
            """.format(
                hasil['legal']['keterlibatan_jaringan'],
                len(hasil['legal']['barang_bukti']),
                "Ada" if hasil['legal']['riwayat_pidana'] else "Tidak Ada"
            ), unsafe_allow_html=True)
        
        with col_hasil3:
            st.markdown("""
            <div class="success-box">
            <strong>📋 REKOMENDASI</strong><br/>
            <strong style="font-size: 1.2em;">{}</strong><br/>
            • Durasi: {}<br/>
            • Tempat: {}
            </div>
            """.format(
                hasil['recommendation']['rekomendasi'],
                hasil['recommendation']['durasi'],
                hasil['recommendation']['tempat'][:50] + "..." if len(hasil['recommendation']['tempat']) > 50 else hasil['recommendation']['tempat']
            ), unsafe_allow_html=True)
        
        st.markdown("---")
        st.subheader("C. DOWNLOAD SURAT HASIL TAT")
        
        st.info("📥 **Surat akan di-generate dalam format Word (.docx) dan PDF**")
        
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            try:
                word_buffer = generate_word_document(
                    hasil['data'],
                    hasil['medical'],
                    hasil['legal'],
                    hasil['recommendation']
                )
                
                filename_word = f"Surat_TAT_{hasil['data']['nama'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                
                st.download_button(
                    label="📘 Download Surat (Word)",
                    data=word_buffer,
                    file_name=filename_word,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Error generating Word: {str(e)}")
        
        with col_dl2:
            try:
                pdf_buffer = generate_pdf_document(
                    hasil['data'],
                    hasil['medical'],
                    hasil['legal'],
                    hasil['recommendation']
                )
                
                filename_pdf = f"Surat_TAT_{hasil['data']['nama'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
                
                st.download_button(
                    label="📕 Download Surat (PDF)",
                    data=pdf_buffer,
                    file_name=filename_pdf,
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Error generating PDF: {str(e)}")
        
        st.markdown("---")
        
        # Preview Surat
        with st.expander("👁️ Preview Isi Surat", expanded=False):
            st.markdown(f"""
            **KEPADA:** {hasil['data']['penerima_surat']}
            
            **PERIHAL:** Hasil Asesmen Terpadu Tersangka a.n. {hasil['data']['nama']}
            
            ---
            
            **IDENTITAS:**
            - Nama: {hasil['data']['nama']}
            - NIK: {hasil['data']['nik']}
            - Tempat/Tgl Lahir: {hasil['data']['tempat_lahir']}, {hasil['data']['tanggal_lahir']}
            - Jenis Kelamin: {hasil['data']['jenis_kelamin']}
            - Alamat: {hasil['data']['alamat']}
            
            ---
            
            **KESIMPULAN:**
            
            a. Bahwa tersangka/terdakwa a.n. **{hasil['data']['nama']}** merupakan penyalahguna 
            narkotika golongan I yaitu **{', '.join(hasil['data']['jenis_narkotika_positif'])}** 
            untuk diri sendiri dengan pola pemakaian **{hasil['medical']['pola_penggunaan']}** 
            kategori **{hasil['medical']['severity_level']}**, didiagnosis 
            **{hasil['medical']['diagnosis']}** ({hasil['medical']['diagnosis_code']}). 
            
            b. Bahwa tersangka/terdakwa a.n. **{hasil['data']['nama']}** 
            **{hasil['legal']['keterlibatan_jaringan']}** indikasi keterlibatan dalam 
            jaringan peredaran gelap narkotika.
            
            ---
            
            **REKOMENDASI:**
            
            **{hasil['recommendation']['rekomendasi']}**
            
            - Durasi: {hasil['recommendation']['durasi']}
            - Tempat: {hasil['recommendation']['tempat']}
            - Tindak Lanjut: {hasil['recommendation']['tindak_lanjut']}
            
            ---
            
            Ditandatangani oleh:
            
            **{hasil['data']['jabatan_penandatangan']}**
            
            **{hasil['data']['nama_penandatangan']}**  
            NIP. {hasil['data']['nip_penandatangan']}
            """)
    
    # =============================================================================
    # TAB 4: PANDUAN
    # (DIPINDAHKAN KE SINI AGAR TAB4 TERDEFINISI DALAM SCOPE main())
    # =============================================================================
    with tab4:
        st.header("📚 IV. PANDUAN PENGGUNAAN SISTEM")
        
        st.markdown("""
        ### 📖 Tentang Sistem
        
        Sistem Asesmen Terpadu (TAT) BNN Provinsi Kalimantan Utara ini dirancang untuk 
        membantu Tim Asesmen Terpadu dalam melakukan asesmen terhadap tersangka/terdakwa 
        penyalahguna narkotika sesuai dengan:
        
        - **UU No. 35 Tahun 2009** tentang Narkotika
        - **KEP/99 I/X/KA/PB/06.00/2025/BNN** tentang Petunjuk Teknis Pelaksanaan Asesmen Terpadu
        - Instrumen Asesmen: ASI, ASAM, DSM-5, ICD-10, PPDGJ III
        """)
        
        st.markdown("---")
        
        with st.expander("🔍 CARA PENGGUNAAN", expanded=True):
            st.markdown("""
            **Langkah 1: Input Data Demografi & Hukum (Tab I)**
            1. Isi data identitas tersangka/terdakwa secara lengkap
            2. Uraikan kronologi kejadian dengan detail
            3. Input data penggunaan narkotika dan hasil tes urine
            4. Isi riwayat hukum (pidana, penahanan, persidangan)
            5. Lengkapi data keterlibatan jaringan dan fakta hukum
            
            **Langkah 2: Asesmen Medis (Tab II)**
            1. Centang kriteria DSM-5 yang terpenuhi
            2. Pilih diagnosis ICD-10 yang sesuai
            3. Isi pola penggunaan narkotika
            4. Lengkapi ASAM 6 dimensi (simplified)
            5. Tulis kesimpulan asesmen medis
            
            **Langkah 3: Generate Surat (Tab III)**
            1. Lengkapi informasi surat (nomor, tanggal, penerima)
            2. Klik tombol "PROSES ASESMEN & GENERATE SURAT"
            3. Review ringkasan hasil
            4. Download surat dalam format Word dan/atau PDF
            """)
        
        with st.expander("📋 KRITERIA DSM-5", expanded=False):
            st.markdown("""
            **11 Kriteria Gangguan Penggunaan Zat (DSM-5):**
            
            **Interpretasi:**
            - **0-1 kriteria**: Tidak ada gangguan
            - **2-3 kriteria**: Gangguan Penggunaan **RINGAN** (Mild)
            - **4-5 kriteria**: Gangguan Penggunaan **SEDANG** (Moderate)  
            - **6-11 kriteria**: Gangguan Penggunaan **BERAT** (Severe)
            """)
            
            for i, crit in enumerate(DSM5_CRITERIA, 1):
                st.markdown(f"{i}. {crit}")
        
        with st.expander("🏥 DIAGNOSIS ICD-10", expanded=False):
            st.markdown("**Kode Diagnosis Gangguan Mental dan Perilaku akibat Penggunaan Zat:**")
            
            for code, diagnosis in DIAGNOSIS_ICD10.items():
                st.markdown(f"- **{code}**: {diagnosis}")
        
        with st.expander("📊 ASAM 6 DIMENSI", expanded=False):
            st.markdown("""
            **American Society of Addiction Medicine (ASAM) Criteria:**
            
            1. **Dimensi 1**: Intoksikasi Akut dan/atau Potensi Withdrawal
               - Menilai tingkat keparahan sakau dan risiko komplikasi medis
            
            2. **Dimensi 2**: Kondisi dan Komplikasi Biomedis
               - Penyakit fisik yang menyertai (HIV, Hepatitis, TBC, dll)
            
            3. **Dimensi 3**: Kondisi Emosional, Behavioral, Kognitif
               - Gangguan mental komorbid (depresi, anxietas, psikotik, dll)
            
            4. **Dimensi 4**: Kesiapan untuk Berubah
               - Motivasi dan insight terhadap masalah ketergantungan
            
            5. **Dimensi 5**: Potensi Relapse, Continued Use, Continued Problem
               - Riwayat relapse, trigger, pola kambuh
            
            6. **Dimensi 6**: Lingkungan Pemulihan/Recovery Environment
               - Dukungan keluarga, kondisi rumah, pekerjaan, ekonomi
            """)
        
        with st.expander("⚖️ DASAR HUKUM", expanded=False):
            st.markdown("""
            **Regulasi yang Menjadi Acuan:**
            
            1. **Undang-Undang Nomor 35 Tahun 2009** tentang Narkotika
               - Pasal 54: Pecandu dan korban wajib menjalani rehabilitasi
               - Pasal 103: Hakim dapat menetapkan rehabilitasi
               - Pasal 127: Penyalahguna dapat direhabilitasi
            
            2. **Peraturan Bersama 7 Instansi Nomor 1 Tahun 2014**
               - Penanganan Pecandu dan Korban Penyalahgunaan Narkotika
               - Mekanisme Asesmen Terpadu
            
            3. **Keputusan BNN Nomor KEP/99 I/X/KA/PB/06.00/2025/BNN**
               - Petunjuk Teknis Pelaksanaan Asesmen Terpadu
               - Format dan prosedur asesmen
            """)
        
        with st.expander("❓ FAQ (Pertanyaan Umum)", expanded=False):
            st.markdown("""
            **Q: Apakah sistem ini menggantikan Tim Asesmen Terpadu?**  
            A: Tidak. Sistem ini adalah **alat bantu** untuk mempermudah proses asesmen. 
            Keputusan final tetap berada di tangan Tim Asesmen Terpadu yang terdiri dari 
            profesional (dokter, psikolog, pekerja sosial, penegak hukum).
            
            **Q: Bagaimana jika data tidak lengkap?**  
            A: Sistem akan memberikan peringatan untuk data wajib yang belum diisi. 
            Isi data selengkap mungkin untuk hasil asesmen yang akurat.
            
            **Q: Apakah bisa mengubah surat setelah di-generate?**  
            A: Surat yang di-download dalam format Word (.docx) dapat diedit secara manual 
            jika diperlukan penyesuaian.
            
            **Q: Bagaimana cara menyimpan data asesmen?**  
            A: Saat ini sistem belum memiliki fitur penyimpanan database. Simpan surat hasil 
            (Word/PDF) sebagai dokumentasi. Pengembangan fitur database akan dilakukan di 
            versi mendatang.
            
            **Q: Apakah rekomendasi sistem pasti tepat?**  
            A: Rekomendasi sistem berdasarkan algoritma rule-based sesuai regulasi. Namun, 
            Tim Asesmen Terpadu tetap harus mempertimbangkan faktor kontekstual lain yang 
            tidak tercakup dalam sistem.
            """)
        
        st.markdown("---")
        
        st.markdown("""
        <div class="warning-box">
        <strong>⚠️ DISCLAIMER:</strong><br/>
        Sistem ini adalah alat bantu untuk proses asesmen. Keputusan final tetap berada 
        di tangan Tim Asesmen Terpadu BNN dan aparat penegak hukum yang berwenang sesuai 
        dengan peraturan perundang-undangan yang berlaku.
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        st.info("""
        **📞 Kontak:**  
        BNN Provinsi Kalimantan Utara  
        Jl. Teuku Umar No. 31, Kota Tarakan, Provinsi Kalimantan Utara  
        """)

if __name__ == "__main__":
    main()
